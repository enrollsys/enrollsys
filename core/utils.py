import os
from django.conf import settings
from django.http import HttpResponse
from docx import Document
from openpyxl import Workbook


def generate_docx(title, content, filename='document.docx'):
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    output_path = os.path.join(settings.MEDIA_ROOT, 'generated', filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_xlsx(headers, rows, filename='data.xlsx'):
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    output_path = os.path.join(settings.MEDIA_ROOT, 'generated', filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    return output_path


def serve_generated_file(filepath):
    if not os.path.exists(filepath):
        return None
    with open(filepath, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(filepath)}"'
        return response

